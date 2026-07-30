from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .render_common import (
    BODY_FONT,
    HEADING_FONT,
    TEST_WARNING,
    TEST_WARNING_ASCII,
    THEME,
    output_file_name,
)


def _set_run_font(run, font: str, size: float, *, bold: bool = False) -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top: int = 70, start: int = 90,
                      bottom: int = 70, end: int = 90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top), ("start", start), ("bottom", bottom), ("end", end)
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _set_fixed_table_layout(table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def _configure_document(document: Document, content: dict, master: dict) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.25

    for name, size in (("Title", 16), ("Heading 1", 12), ("Heading 2", 11)):
        style = document.styles[name]
        style.font.name = HEADING_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(THEME["navy"])

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(
        f"{master.get('project_name', '工程项目')}｜"
        f"{content.get('logical_id', '')}"
    )
    _set_run_font(run, BODY_FONT, 8)
    run.font.color.rgb = RGBColor.from_string(THEME["slate"])

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    warning = footer.add_run(
        TEST_WARNING + "　|　" + TEST_WARNING_ASCII + "　|　第"
    )
    _set_run_font(warning, BODY_FONT, 8, bold=True)
    warning.font.color.rgb = RGBColor.from_string(THEME["alert"])
    _add_page_number(footer)
    tail = footer.add_run("页")
    _set_run_font(tail, BODY_FONT, 8)

    document.core_properties.title = content.get("title", "")
    document.core_properties.subject = TEST_WARNING
    document.core_properties.author = "TEST-资料编制系统"
    document.core_properties.comments = TEST_WARNING


def _add_metadata_table(document: Document, content: dict) -> None:
    rows = [
        ("文件编号", content.get("document_number", content.get("logical_id", ""))),
        ("版本／日期", f"{content.get('revision', 'A')}／{content.get('date', '')}"),
        ("资料属性", TEST_WARNING),
    ]
    table = document.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _set_fixed_table_layout(table)
    for row_index, (key, value) in enumerate(rows):
        table.cell(row_index, 0).width = Cm(3.2)
        table.cell(row_index, 1).width = Cm(12.8)
        table.cell(row_index, 0).text = key
        table.cell(row_index, 1).text = str(value)
        _shade(table.cell(row_index, 0), THEME["pale"])
        for cell in table.rows[row_index].cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    _set_run_font(run, BODY_FONT, 9, bold=(cell is table.cell(row_index, 0)))
    document.add_paragraph()


def _add_business_table(document: Document, table_data: dict) -> None:
    title = table_data.get("title")
    if title:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(title)
        _set_run_font(run, HEADING_FONT, 10.5, bold=True)
        run.font.color.rgb = RGBColor.from_string(THEME["navy"])
    headers = [str(value) for value in table_data.get("headers", [])]
    rows = table_data.get("rows", [])
    if not headers:
        return
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        _shade(cell, THEME["navy"])
        for run in cell.paragraphs[0].runs:
            _set_run_font(run, BODY_FONT, 8.5, bold=True)
            run.font.color.rgb = RGBColor(255, 255, 255)
    for data_row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(data_row):
            if index >= len(cells):
                break
            cells[index].text = "" if value is None else str(value)
            if "不合格" in str(value) or "异常" in str(value):
                _shade(cells[index], THEME["exception"])
            elif "合格" in str(value) or "闭环" in str(value):
                _shade(cells[index], THEME["qualified"])
            for paragraph in cells[index].paragraphs:
                for run in paragraph.runs:
                    _set_run_font(run, BODY_FONT, 8.5)
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
def _add_approvals(document: Document, approvals: Iterable[dict]) -> None:
    approvals = list(approvals)
    if not approvals:
        return
    paragraph = document.add_paragraph()
    run = paragraph.add_run("编审记录")
    _set_run_font(run, HEADING_FONT, 10.5, bold=True)
    _add_business_table(
        document,
        {
            "headers": ["角色", "人员", "日期", "记录方式"],
            "rows": [
                [
                    row.get("role", ""),
                    row.get("name", ""),
                    row.get("date", ""),
                    row.get("record", "电子记录（测试）"),
                ]
                for row in approvals
            ],
        },
    )


def _add_evidence_panel(document: Document, panel: dict) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(panel.get("label", "测试模拟证据面板"))
    _set_run_font(run, HEADING_FONT, 10.5, bold=True)
    run.font.color.rgb = RGBColor.from_string(THEME["alert"])
    table = document.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    top, bottom = table.cell(0, 0), table.cell(1, 0)
    _shade(top, "18212B")
    _shade(bottom, "2C3845")
    top.text = panel.get(
        "pattern", "█▓▒░  ───────  ░▒▓██▓▒░  ───────  ░▒▓█"
    )
    bottom.text = (
        f"对象：{panel.get('object', '')}　"
        f"{panel.get('annotation', '')}"
    )
    for cell, size in ((top, 15), (bottom, 9)):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_margins(cell, top=120, bottom=120)
        for cell_paragraph in cell.paragraphs:
            cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for cell_run in cell_paragraph.runs:
                _set_run_font(cell_run, BODY_FONT, size, bold=True)
                cell_run.font.color.rgb = RGBColor(255, 255, 255)


def render_docx(content: dict, master: dict, output: Path) -> Path:
    output.mkdir(parents=True, exist_ok=True)
    document = Document()
    _configure_document(document, content, master)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(content.get("title", "工程测试资料"))
    _set_run_font(run, HEADING_FONT, 16, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("业务完整性验收测试资料")
    _set_run_font(run, BODY_FONT, 10)
    run.font.color.rgb = RGBColor.from_string(THEME["slate"])

    _add_metadata_table(document, content)

    for section in content.get("sections", []):
        heading = document.add_paragraph(style="Heading 1")
        run = heading.add_run(section.get("heading", ""))
        _set_run_font(run, HEADING_FONT, 12, bold=True)
        for text in section.get("paragraphs", []):
            paragraph = document.add_paragraph(str(text))
            paragraph.paragraph_format.first_line_indent = Cm(0.74)
        if section.get("bullets"):
            for bullet in section["bullets"]:
                paragraph = document.add_paragraph(style="List Bullet")
                run = paragraph.add_run(str(bullet))
                _set_run_font(run, BODY_FONT, 10.5)
        if section.get("key_values"):
            _add_business_table(
                document,
                {
                    "headers": ["项目", "内容"],
                    "rows": [[key, value] for key, value in section["key_values"]],
                },
            )

    for table_data in content.get("tables", []):
        _add_business_table(document, table_data)

    for panel in content.get("evidence_panels", []):
        _add_evidence_panel(document, panel)

    if content.get("references"):
        heading = document.add_paragraph(style="Heading 1")
        run = heading.add_run("引用文件")
        _set_run_font(run, HEADING_FONT, 12, bold=True)
        for reference in content["references"]:
            paragraph = document.add_paragraph(style="List Bullet")
            run = paragraph.add_run(str(reference))
            _set_run_font(run, BODY_FONT, 9.5)

    _add_approvals(document, content.get("approvals", []))

    path = output / output_file_name(content, "docx")
    document.save(path)
    return path
